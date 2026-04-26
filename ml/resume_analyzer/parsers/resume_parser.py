"""
resume_parser.py
----------------
Phase 2 — Robust Resume Ingestion & Pre-processing

Handles:
  - PDF (single-column, multi-column, scanned detection)
  - DOCX
  - Smart text normalization (preserves C++, .NET, C#)
  - Section detection (Education, Experience, Skills, etc.)
  - Outputs a clean structured dict for downstream NLP
"""

import re
import io
import logging
import datetime
from pathlib import Path
from typing import Optional

# PDF parsing
import pdfplumber
from sklearn.cluster import KMeans
import numpy as np

# DOCX parsing
from docx import Document as DocxDocument
from docx.oxml.ns import qn

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Section heading patterns — order matters (most specific first)
# ---------------------------------------------------------------------------
SECTION_PATTERNS = {
    "contact":     r"(contact|personal\s+info|personal\s+details|reach\s+me)",
    "summary":     r"(summary|objective|profile|about\s+me|career\s+objective|professional\s+summary)",
    "education":   r"(education|academic|qualification|degree|study|studies)",
    "experience":  r"(experience|employment|work\s+history|career|internship|job)",
    "skills":      r"(skills|technical\s+skills|competencies|technologies|tools|proficiencies)",
    "projects":    r"(projects|portfolio|work\s+samples|personal\s+projects|academic\s+projects)",
    "certifications": r"(certification|certificate|licenses?|accreditation|courses?)",
    "awards":      r"(awards?|honors?|achievements?|recognition|accomplishments?)",
    "publications": r"(publications?|papers?|research|journal)",
    "languages":   r"(languages?|spoken\s+languages?|linguistic)",
    "references":  r"(references?|referees?)",
    "volunteer":   r"(volunteer|community|social\s+work|extra.?curricular)",
}

# Technical terms whose punctuation must be preserved
PRESERVE_TECHNICAL = [
    r"C\+\+", r"C#", r"\.NET", r"ASP\.NET", r"F#", r"R\b",
    r"Node\.js", r"React\.js", r"Vue\.js", r"Next\.js", r"Express\.js",
    r"Three\.js", r"D3\.js", r"\bCI/CD\b", r"REST\b", r"GraphQL",
    r"NoSQL", r"HTML5", r"CSS3", r"ES6\+", r"TypeScript",
    r"TensorFlow", r"PyTorch", r"scikit-learn", r"OpenCV",
]


# ===========================================================================
# PDF Parser
# ===========================================================================

class PDFParser:
    """
    Extracts text from PDFs using pdfplumber with KMeans clustering
    to natively solve the two-column bleed-through issue.
    """

    def __init__(self, column_threshold: float = 0.25):
        self.column_threshold = column_threshold

    def is_scanned(self, file_bytes: bytes) -> bool:
        """
        Returns True if the PDF appears to be a scanned image
        (no selectable text layer found).
        """
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                total_chars = 0
                for page in pdf.pages[:3]:
                    text = page.extract_text()
                    if text:
                        total_chars += len(text.strip())
                return total_chars < 50
        except Exception as e:
            logger.error(f"Error checking scanned PDF: {e}")
            return True

    def _group_words_into_lines(self, words: list) -> list:
        # Sort words by top coordinate, then x0
        words.sort(key=lambda w: (w['top'], w['x0']))
        lines = []
        if not words:
            return lines
            
        current_line = {
            'text': words[0]['text'],
            'x0': float(words[0]['x0']),
            'x1': float(words[0]['x1']),
            'top': float(words[0]['top']),
            'bottom': float(words[0]['bottom'])
        }
        
        for w in words[1:]:
            # If word is on the same line (y difference is small)
            if abs(w['top'] - current_line['top']) < 5:
                # Add a space between words
                current_line['text'] += " " + w['text']
                current_line['bottom'] = max(current_line['bottom'], float(w['bottom']))
                current_line['x1'] = max(current_line['x1'], float(w['x1']))
            else:
                lines.append(current_line)
                current_line = {
                    'text': w['text'],
                    'x0': float(w['x0']),
                    'x1': float(w['x1']),
                    'top': float(w['top']),
                    'bottom': float(w['bottom'])
                }
        lines.append(current_line)
        return lines

    def extract_text(self, file_bytes: bytes) -> str:
        """
        Main extraction method using KMeans column detection.
        """
        all_pages_text = []

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                # Extract words with bounding boxes
                words = page.extract_words(x_tolerance=3, y_tolerance=3, keep_blank_chars=True)
                if not words:
                    continue
                    
                lines = self._group_words_into_lines(words)
                if not lines:
                    continue
                    
                x0_coords = np.array([line['x0'] for line in lines]).reshape(-1, 1)
                
                page_width = float(page.width)
                range_x0 = x0_coords.max() - x0_coords.min()
                
                # Determine if page is multi-column based on x0 variance/range
                if range_x0 > page_width * self.column_threshold and len(lines) > 5:
                    # 1D KMeans to find 2 columns
                    kmeans = KMeans(n_clusters=2, random_state=42, n_init=10)
                    labels = kmeans.fit_predict(x0_coords)
                    centers = kmeans.cluster_centers_.flatten()
                    
                    # Ensure cluster 0 is left, cluster 1 is right
                    if centers[0] > centers[1]:
                        labels = 1 - labels
                        
                    left_lines = [lines[i] for i in range(len(lines)) if labels[i] == 0]
                    right_lines = [lines[i] for i in range(len(lines)) if labels[i] == 1]
                    
                    # Separate out headers that span the whole width (x0 is left, but x1 is far right)
                    headers = []
                    pure_left = []
                    
                    for l in left_lines:
                        if l['x1'] > page_width * 0.75:
                            headers.append(l)
                        else:
                            pure_left.append(l)
                            
                    # Sort top-to-bottom within each group
                    headers.sort(key=lambda l: l['top'])
                    pure_left.sort(key=lambda l: l['top'])
                    right_lines.sort(key=lambda l: l['top'])
                    
                    # Merge logic: Headers -> Left -> Right
                    page_text_parts = [l['text'] for l in headers] + \
                                      [l['text'] for l in pure_left] + \
                                      [l['text'] for l in right_lines]
                    page_text = "\n".join(page_text_parts)
                else:
                    # Single column
                    lines.sort(key=lambda l: l['top'])
                    page_text = "\n".join([l['text'] for l in lines])
                    
                all_pages_text.append(page_text)

        return "\n\n".join(all_pages_text)


# ===========================================================================
# DOCX Parser
# ===========================================================================

class DOCXParser:
    """
    Extracts text from .docx files including text inside tables,
    which standard python-docx iteration misses.
    """

    def extract_text(self, file_bytes: bytes) -> str:
        doc = DocxDocument(io.BytesIO(file_bytes))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Extract text from tables (many resumes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append(" | ".join(row_texts))

        return "\n".join(parts)


# ===========================================================================
# Text Normalizer
# ===========================================================================

class TextNormalizer:
    """
    Cleans and normalizes resume text while preserving technical terms.
    """

    def __init__(self):
        # Build a placeholder map to protect technical terms during normalization
        self._placeholders = {}

    def normalize(self, text: str) -> str:
        text = self._protect_technical_terms(text)
        text = self._fix_encoding_artifacts(text)
        text = self._normalize_whitespace(text)
        text = self._restore_technical_terms(text)
        return text.strip()

    def _protect_technical_terms(self, text: str) -> str:
        """Replace technical terms with safe placeholders before any cleaning."""
        self._placeholders = {}
        for i, pattern in enumerate(PRESERVE_TECHNICAL):
            placeholder = f"__TECH_{i}__"
            matches = re.findall(pattern, text, flags=re.IGNORECASE)
            for match in set(matches):
                self._placeholders[placeholder] = match
                text = text.replace(match, placeholder)
        return text

    def _restore_technical_terms(self, text: str) -> str:
        for placeholder, original in self._placeholders.items():
            text = text.replace(placeholder, original)
        return text

    def _fix_encoding_artifacts(self, text: str) -> str:
        # Fix common PDF encoding issues
        replacements = {
            "\ufb01": "fi", "\ufb02": "fl", "\u2013": "-",
            "\u2014": "-", "\u2019": "'", "\u201c": '"',
            "\u201d": '"', "\u00a0": " ", "\x00": "",
        }
        for bad, good in replacements.items():
            text = text.replace(bad, good)
        return text

    def _normalize_whitespace(self, text: str) -> str:
        # Collapse multiple spaces but preserve newlines
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text


# ===========================================================================
# Section Detector
# ===========================================================================

class SectionDetector:
    """
    Identifies and splits the resume into named sections.
    Uses heuristics: short lines in ALL CAPS or Title Case at section boundaries.
    """

    def detect_sections(self, text: str) -> dict:
        lines = text.split("\n")
        sections = {key: [] for key in SECTION_PATTERNS}
        sections["other"] = []

        current_section = "other"

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            detected = self._classify_line(stripped)
            if detected:
                current_section = detected
            else:
                sections[current_section].append(stripped)

        # Convert lists to strings
        return {k: "\n".join(v).strip() for k, v in sections.items()}

    def _classify_line(self, line: str) -> Optional[str]:
        """
        Returns section key if this line looks like a section heading, else None.
        Heuristics:
          - Short line (< 60 chars)
          - ALL CAPS or Title Case
          - Matches one of our section patterns
        """
        if len(line) > 60:
            return None

        is_heading_case = line.isupper() or line.istitle() or (
            line.replace(" ", "").replace("-", "").replace("/", "").isalpha()
            and line[0].isupper()
        )

        if not is_heading_case:
            return None

        lower = line.lower()
        for section_key, pattern in SECTION_PATTERNS.items():
            if re.search(pattern, lower):
                return section_key

        return None


# ===========================================================================
# Main ResumeParser — orchestrates everything
# ===========================================================================

class ResumeParser:
    """
    Main entry point.

    Usage:
        parser = ResumeParser()
        result = parser.parse(file_bytes, filename="resume.pdf")
    """

    def __init__(self):
        self.pdf_parser = PDFParser()
        self.docx_parser = DOCXParser()
        self.normalizer = TextNormalizer()
        self.section_detector = SectionDetector()

    def parse(self, file_bytes: bytes, filename: str) -> dict:
        """
        Parse a resume file and return structured data.

        Returns:
        {
            "filename": str,
            "file_type": "pdf" | "docx" | "unknown",
            "is_scanned": bool,           # True = image PDF, ATS will penalize
            "raw_text": str,              # normalized full text
            "sections": {
                "summary": str,
                "education": str,
                "experience": str,
                "skills": str,
                "projects": str,
                "certifications": str,
                ... (other detected sections)
            },
            "word_count": int,
            "has_contact_info": bool,
            "warnings": [str]             # parser-level warnings
        }
        """
        filename_lower = filename.lower()
        warnings = []
        is_scanned = False

        # --- Step 1: Extract raw text based on file type ---
        if filename_lower.endswith(".pdf"):
            file_type = "pdf"
            is_scanned = self.pdf_parser.is_scanned(file_bytes)
            if is_scanned:
                warnings.append(
                    "⚠️ This PDF appears to be a scanned image. "
                    "No text layer detected — ATS systems cannot read it. "
                    "Please upload a text-based PDF."
                )
                raw_text = ""
            else:
                raw_text = self.pdf_parser.extract_text(file_bytes)

        elif filename_lower.endswith(".docx"):
            file_type = "docx"
            raw_text = self.docx_parser.extract_text(file_bytes)

        elif filename_lower.endswith(".doc"):
            file_type = "doc"
            warnings.append(
                "⚠️ Old .doc format detected. Please save as .docx or .pdf for best results."
            )
            raw_text = ""

        else:
            file_type = "unknown"
            warnings.append("⚠️ Unsupported file format. Please upload PDF or DOCX.")
            raw_text = ""
            
        if not raw_text.strip():
            raise ValueError(f"Could not extract any text from {filename}. The file may be empty, an unsupported format (like older .doc), or a scanned image.")

        # --- Step 2: Normalize text ---
        normalized_text = self.normalizer.normalize(raw_text)

        # --- Step 3: Detect sections ---
        sections = self.section_detector.detect_sections(normalized_text)

        # --- Step 4: Check for contact info ---
        has_contact = self._detect_contact_info(normalized_text)
        if not has_contact:
            warnings.append("⚠️ No contact information detected (email/phone).")

        # --- Step 5: Completeness warnings ---
        important_sections = ["experience", "education", "skills"]
        for sec in important_sections:
            if not sections.get(sec):
                warnings.append(
                    f"⚠️ Section '{sec.capitalize()}' not found or empty. "
                    f"Most ATS systems require this section."
                )

        word_count = len(normalized_text.split())

        if word_count < 150:
            warnings.append("⚠️ Resume seems too short (under 150 words). Consider adding more detail.")
        elif word_count > 1200:
            warnings.append("⚠️ Resume is very long (over 1200 words). Consider condensing to 1-2 pages.")

        logger.info(
            f"Parsed '{filename}' | type={file_type} | "
            f"words={word_count} | scanned={is_scanned} | "
            f"sections_found={[k for k,v in sections.items() if v]}"
        )

        return {
            "filename": filename,
            "file_type": file_type,
            "is_scanned": is_scanned,
            "raw_text": normalized_text,
            "sections": sections,
            "word_count": word_count,
            "has_contact_info": has_contact,
            "warnings": warnings,
        }

    def _detect_contact_info(self, text: str) -> bool:
        email_pattern = r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+"
        phone_pattern = r"(\+?\d[\d\s\-().]{7,}\d)"
        has_email = bool(re.search(email_pattern, text))
        has_phone = bool(re.search(phone_pattern, text))
        return has_email or has_phone
